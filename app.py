from flask import Flask, render_template, request, jsonify, send_file
from docx import Document
import os

app = Flask(__name__)

@app.route('/')
def index():
    return render_template('index.html')

# Generate and return a preview of the modified document
@app.route('/preview', methods=['POST'])
def preview():
    # Get user input
    name = request.form.get('name')
    email = request.form.get('email')
    message = request.form.get('message')

    # Load and modify the .docx template
    doc = Document('template.docx')

    # Replace placeholders in the document
    for para in doc.paragraphs:
        if '[NAME]' in para.text:
            para.text = para.text.replace('[NAME]', name)
        if '[EMAIL]' in para.text:
            para.text = para.text.replace('[EMAIL]', email)
        if '[MESSAGE]' in para.text:
            para.text = para.text.replace('[MESSAGE]', message)

    # Convert the document content to plain text for preview
    preview_text = ''
    for para in doc.paragraphs:
        preview_text += para.text + '<br>'

    # Return the text as HTML for preview
    return jsonify({'preview': preview_text})

# Generate and download the modified docx file
@app.route('/generate', methods=['POST'])
def generate():
    # Get user input
    name = request.form.get('name')
    email = request.form.get('email')
    message = request.form.get('message')

    # Load and modify the .docx template
    doc = Document('template.docx')

    # Replace placeholders in the document
    for para in doc.paragraphs:
        if '[NAME]' in para.text:
            para.text = para.text.replace('[NAME]', name)
        if '[EMAIL]' in para.text:
            para.text = para.text.replace('[EMAIL]', email)
        if '[MESSAGE]' in para.text:
            para.text = para.text.replace('[MESSAGE]', message)

    # Save the modified docx to a temporary location
    output_path = 'modified_template.docx'
    doc.save(output_path)

    # Send the modified file back to the user for download
    return send_file(output_path, as_attachment=True, download_name='modified_document.docx')

if __name__ == '__main__':
    app.run(debug=True)
